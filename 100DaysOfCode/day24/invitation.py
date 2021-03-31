"""pseudo-code:
1. read from a list of people
    iterate over it 
2. replace a specific space holder within a word doc with a name of a person in the list
3. save as a new word document with person name
 

 """

from docx import Document
import csv 

with open()
people_invited= ["Maureen","Ahmed"]

# new document 
document = Document()

document.add_heading('Wedding invitation')

for name in people_invited:
    paragraph = document.add_paragraph('We are')
